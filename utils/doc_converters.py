"""utils/doc_converters.py

Lightweight Markdown → PDF and Markdown → DOCX converters used by the
"My Paper Drafts" view.

The MD subset supported is the one needed by ``MPC_Template.md``:

- ATX headings (``#`` … ``####``)
- Paragraphs (separated by blank lines)
- Unordered lists (``- `` / ``* ``)
- Ordered lists (``1. ``)
- Block quotes (``> ``)
- Horizontal rules (``---`` on their own line)
- GitHub-style pipe tables with a header separator row
- Inline ``**bold**``, ``*italic*``, ``` `code` ```

Anything outside this subset is rendered as plain text. We intentionally avoid
``weasyprint``/``pypandoc`` so the converter stays Streamlit-Cloud friendly.
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Iterator

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether,
)


# ─── Block parser ────────────────────────────────────────────────────────────

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_HR_RE = re.compile(r"^\s*---+\s*$")
_BULLET_RE = re.compile(r"^[\-\*]\s+(.*)$")
_ORDERED_RE = re.compile(r"^\d+\.\s+(.*)$")
_QUOTE_RE = re.compile(r"^>\s?(.*)$")


def _parse_blocks(md: str) -> list[dict]:
    """Tokenise the markdown into a flat list of block dicts.

    Each block is one of:
      {type: "heading", level: int, text: str}
      {type: "paragraph", text: str}
      {type: "ulist", items: list[str]}
      {type: "olist", items: list[str]}
      {type: "quote", text: str}
      {type: "hr"}
      {type: "table", header: list[str], rows: list[list[str]]}
    """
    lines = md.splitlines()
    blocks: list[dict] = []
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if _HR_RE.match(line):
            blocks.append({"type": "hr"})
            i += 1
            continue

        m = _HEADING_RE.match(line)
        if m:
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2)})
            i += 1
            continue

        # Pipe table: a row that starts with `|` and the next line is a separator
        if stripped.startswith("|") and i + 1 < n and re.match(r"^\s*\|?\s*[:\-\s|]+\|?\s*$", lines[i + 1]):
            header = _split_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(_split_table_row(lines[i].strip()))
                i += 1
            blocks.append({"type": "table", "header": header, "rows": rows})
            continue

        m = _BULLET_RE.match(stripped)
        if m:
            items = [m.group(1)]
            i += 1
            while i < n:
                m2 = _BULLET_RE.match(lines[i].strip())
                if not m2 or not lines[i].strip():
                    break
                items.append(m2.group(1))
                i += 1
            blocks.append({"type": "ulist", "items": items})
            continue

        m = _ORDERED_RE.match(stripped)
        if m:
            items = [m.group(1)]
            i += 1
            while i < n:
                m2 = _ORDERED_RE.match(lines[i].strip())
                if not m2 or not lines[i].strip():
                    break
                items.append(m2.group(1))
                i += 1
            blocks.append({"type": "olist", "items": items})
            continue

        m = _QUOTE_RE.match(line)
        if m:
            buf = [m.group(1)]
            i += 1
            while i < n:
                m2 = _QUOTE_RE.match(lines[i])
                if not m2 or not lines[i].strip():
                    break
                buf.append(m2.group(1))
                i += 1
            blocks.append({"type": "quote", "text": " ".join(buf).strip()})
            continue

        # Paragraph: collect until blank line or block boundary
        buf = [stripped]
        i += 1
        while i < n:
            ln = lines[i]
            if not ln.strip():
                break
            if (
                _HEADING_RE.match(ln)
                or _HR_RE.match(ln)
                or _BULLET_RE.match(ln.strip())
                or _ORDERED_RE.match(ln.strip())
                or ln.strip().startswith("|")
                or _QUOTE_RE.match(ln)
            ):
                break
            buf.append(ln.strip())
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(buf)})

    return blocks


def _split_table_row(row: str) -> list[str]:
    parts = row.strip().strip("|").split("|")
    return [p.strip() for p in parts]


# ─── Inline formatting ───────────────────────────────────────────────────────

# Order matters: code first (consumes its content), then bold, then italic.
_INLINE_CODE_RE = re.compile(r"`([^`]+)`")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def _md_inline_to_reportlab(text: str) -> str:
    """Convert inline MD to ReportLab's HTML-ish Paragraph markup."""
    import html
    out = html.escape(text)
    out = _INLINE_CODE_RE.sub(r'<font face="Courier">\1</font>', out)
    out = _BOLD_RE.sub(r"<b>\1</b>", out)
    out = _ITALIC_RE.sub(r"<i>\1</i>", out)
    return out


def _strip_md_inline(text: str) -> Iterator[tuple[str, dict]]:
    """Yield (chunk, style_flags) tuples for python-docx, e.g. {'bold': True}."""
    pos = 0
    pattern = re.compile(
        r"(\*\*(?P<b>.+?)\*\*)|(?<!\*)\*(?!\*)(?P<i>.+?)(?<!\*)\*(?!\*)|`(?P<c>[^`]+)`"
    )
    for m in pattern.finditer(text):
        if m.start() > pos:
            yield text[pos:m.start()], {}
        if m.group("b") is not None:
            yield m.group("b"), {"bold": True}
        elif m.group("i") is not None:
            yield m.group("i"), {"italic": True}
        elif m.group("c") is not None:
            yield m.group("c"), {"code": True}
        pos = m.end()
    if pos < len(text):
        yield text[pos:], {}


# ─── Markdown → PDF ──────────────────────────────────────────────────────────

def md_to_pdf(md: str, title: str | None = None) -> BytesIO:
    """Render a markdown document to a PDF (BytesIO ready for download)."""
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=title or "Paper Draft",
    )
    styles = getSampleStyleSheet()

    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10, leading=14, spaceAfter=6)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=16, leading=20, spaceBefore=10, spaceAfter=8)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=13, leading=17, spaceBefore=10, spaceAfter=6)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=11, leading=15, spaceBefore=8, spaceAfter=5)
    h4 = ParagraphStyle("H4", parent=h3, fontSize=10)
    quote = ParagraphStyle(
        "Quote", parent=body, leftIndent=14, textColor=colors.HexColor("#555555"),
        fontName="Helvetica-Oblique", spaceAfter=6,
    )

    story = []
    if title:
        story.append(Paragraph(_md_inline_to_reportlab(title), h1))
        story.append(Spacer(1, 4))

    for block in _parse_blocks(md):
        bt = block["type"]
        if bt == "heading":
            level = block["level"]
            style = {1: h1, 2: h2, 3: h3}.get(level, h4)
            story.append(Paragraph(_md_inline_to_reportlab(block["text"]), style))
        elif bt == "paragraph":
            story.append(Paragraph(_md_inline_to_reportlab(block["text"]), body))
        elif bt == "quote":
            story.append(Paragraph(_md_inline_to_reportlab(block["text"]), quote))
        elif bt == "hr":
            story.append(Spacer(1, 4))
            story.append(HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#cccccc")))
            story.append(Spacer(1, 4))
        elif bt == "ulist":
            for it in block["items"]:
                story.append(Paragraph(f"• {_md_inline_to_reportlab(it)}", body))
        elif bt == "olist":
            for idx, it in enumerate(block["items"], start=1):
                story.append(Paragraph(f"{idx}. {_md_inline_to_reportlab(it)}", body))
        elif bt == "table":
            header = [Paragraph(_md_inline_to_reportlab(c), body) for c in block["header"]]
            rows = [
                [Paragraph(_md_inline_to_reportlab(c), body) for c in r]
                for r in block["rows"]
            ]
            data = [header] + rows
            n_cols = max(len(r) for r in data) if data else 1
            avail = doc.width
            col_w = avail / n_cols
            t = Table(data, colWidths=[col_w] * n_cols, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eeeeee")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#999999")),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#bbbbbb")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(KeepTogether(t))
            story.append(Spacer(1, 6))

    doc.build(story)
    buf.seek(0)
    return buf


# ─── Markdown → DOCX ─────────────────────────────────────────────────────────

_DOCX_HEADING_STYLES = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}


def _docx_add_runs(paragraph, text: str) -> None:
    for chunk, flags in _strip_md_inline(text):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        if flags.get("bold"):
            run.bold = True
        if flags.get("italic"):
            run.italic = True
        if flags.get("code"):
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def md_to_docx(md: str, title: str | None = None) -> BytesIO:
    """Render a markdown document to a DOCX (BytesIO ready for download)."""
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    section = document.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    if title:
        head = document.add_paragraph()
        head.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = head.add_run(title)
        run.bold = True
        run.font.size = Pt(18)

    for block in _parse_blocks(md):
        bt = block["type"]
        if bt == "heading":
            level = block["level"]
            p = document.add_paragraph(style=_DOCX_HEADING_STYLES.get(level, "Heading 4"))
            _docx_add_runs(p, block["text"])
        elif bt == "paragraph":
            p = document.add_paragraph()
            _docx_add_runs(p, block["text"])
        elif bt == "quote":
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.7)
            run_intro = p.add_run("")
            _docx_add_runs(p, block["text"])
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif bt == "hr":
            document.add_paragraph("─" * 60)
        elif bt == "ulist":
            for it in block["items"]:
                p = document.add_paragraph(style="List Bullet")
                _docx_add_runs(p, it)
        elif bt == "olist":
            for it in block["items"]:
                p = document.add_paragraph(style="List Number")
                _docx_add_runs(p, it)
        elif bt == "table":
            header = block["header"]
            rows = block["rows"]
            n_cols = max([len(header)] + [len(r) for r in rows]) if rows else len(header)
            table = document.add_table(rows=1 + len(rows), cols=n_cols)
            table.style = "Light Grid Accent 1"
            # Header
            for j, cell_text in enumerate(header):
                cell = table.cell(0, j)
                p = cell.paragraphs[0]
                _docx_add_runs(p, cell_text)
                for run in p.runs:
                    run.bold = True
            # Body
            for i, row in enumerate(rows, start=1):
                for j in range(n_cols):
                    cell = table.cell(i, j)
                    cell_text = row[j] if j < len(row) else ""
                    p = cell.paragraphs[0]
                    p.text = ""
                    _docx_add_runs(p, cell_text)
            document.add_paragraph()

    buf = BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf
